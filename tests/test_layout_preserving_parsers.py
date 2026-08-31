from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

from app.services.extraction_service import ExtractionService  # noqa: F401
from app.core.parser import compress_text_content
from app.core.parser.doc_parser import _clean_word_text, _extract_table_data
from app.core.parser.pdf_parser import parse_pdf
from app.core.parser.word_parser import parse_word


def test_pdf_uses_layout_extraction_mode(tmp_path: Path):
    pdf_path = tmp_path / "form.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 test")
    page = MagicMock()
    page.extract_text.return_value = "POR      POL      POD\n         NANSHA  SYDNEY"
    page.images = []

    with patch("app.core.parser.pdf_parser.PdfReader") as reader:
        reader.return_value.pages = [page]
        text, tables, ocr_text = parse_pdf(pdf_path)

    page.extract_text.assert_called_once_with(extraction_mode="layout")
    assert "         NANSHA" in text
    assert tables == []
    assert ocr_text == ""


def test_pdf_falls_back_when_layout_mode_is_unavailable(tmp_path: Path):
    pdf_path = tmp_path / "legacy.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 test")
    page = MagicMock()
    page.extract_text.side_effect = [TypeError("unsupported"), "plain text"]
    page.images = []

    with patch("app.core.parser.pdf_parser.PdfReader") as reader:
        reader.return_value.pages = [page]
        text, _, _ = parse_pdf(pdf_path)

    assert text == "plain text"
    assert page.extract_text.call_count == 2


def test_docx_keeps_paragraph_and_table_document_order(tmp_path: Path):
    doc_path = tmp_path / "ordered.docx"
    document = Document()
    document.add_paragraph("BEFORE TABLE")
    table = document.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "POR"
    table.rows[0].cells[1].text = "POL"
    table.rows[0].cells[2].text = "POD"
    table.rows[0].cells[3].text = "FPOD"
    table.rows[1].cells[1].text = "NANSHA"
    table.rows[1].cells[2].text = "SYDNEY"
    table.rows[1].cells[3].text = "SYDNEY"
    document.add_paragraph("AFTER TABLE")
    document.save(doc_path)

    text, tables, ocr_text = parse_word(doc_path)

    assert text.index("BEFORE TABLE") < text.index("| POR | POL | POD | FPOD |")
    assert text.index("|  | NANSHA | SYDNEY | SYDNEY |") < text.index("AFTER TABLE")
    assert tables[0]["rows"][1] == ["", "NANSHA", "SYDNEY", "SYDNEY"]
    assert ocr_text == ""


def test_legacy_doc_keeps_leading_and_trailing_empty_cells():
    # The final tab is Word's row terminator; the preceding one represents a
    # real empty last cell and must survive table parsing.
    cleaned = _clean_word_text("\tNANSHA\tSYDNEY\t\t\r")
    tables = _extract_table_data(cleaned)

    assert cleaned == "\tNANSHA\tSYDNEY\t\t"
    assert tables[0]["rows"][0] == ["", "NANSHA", "SYDNEY", ""]


def test_compression_preserves_layout_and_original_line_order():
    source = (
        "POR          POL          POD          FPOD\n"
        "             NANSHA       SYDNEY       SYDNEY\n"
        + "irrelevant reference row\n" * 100
    )

    compressed = compress_text_content(source, max_chars=180)

    assert compressed.index("POR") < compressed.index("NANSHA")
    value_line = next(line for line in compressed.splitlines() if "NANSHA" in line)
    assert value_line.startswith("             NANSHA")
