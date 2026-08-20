import io
import pytest
from unittest.mock import AsyncMock, MagicMock, patch
from PIL import Image
import docx
from sqlalchemy import delete

from app.config import settings
from app.database import AsyncSessionLocal, init_db
from app.models.system import SystemConfig
from app.services.vision_service import (
    MAX_VISION_INPUT_BYTES,
    MAX_VISION_OUTPUT_CHARS,
    VisionBudget,
    VisionService,
)
from app.core.parser.word_parser import parse_word
from app.core.parser.pdf_parser import parse_pdf
from app.core.parser import parse_single_file


def _create_test_image(width: int, height: int, color=(255, 255, 255), fmt="JPEG") -> bytes:
    img = Image.new("RGB", (width, height), color=color)
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return buf.getvalue()


def test_is_valid_document_image():
    # Large document screenshot
    large_img = _create_test_image(800, 600)
    assert VisionService.is_valid_document_image(large_img) is True

    # Tiny decoration icon (e.g. 50x50)
    tiny_img = _create_test_image(50, 50)
    assert VisionService.is_valid_document_image(tiny_img) is False

    # Empty or corrupt bytes
    assert VisionService.is_valid_document_image(b"") is False
    assert VisionService.is_valid_document_image(b"not_an_image") is False


def test_optimize_image_for_vision():
    # Oversized image (e.g. 3840x2160 4K)
    huge_img = _create_test_image(3840, 2160)
    optimized = VisionService.optimize_image_for_vision(huge_img, max_dim=1920)
    
    with Image.open(io.BytesIO(optimized)) as img:
        assert max(img.size) <= 1920
        assert img.format == "JPEG"


def test_invalid_or_oversized_image_is_never_forwarded():
    assert VisionService.optimize_image_for_vision(b"not-an-image") is None
    assert VisionService.optimize_image_for_vision(b"x" * (MAX_VISION_INPUT_BYTES + 1)) is None

    with patch("httpx.Client.post") as mock_post:
        result = VisionService.transcribe_image_sync(
            b"not-an-image",
            enabled=True,
            custom_api_key="mock_test_key",
        )
    assert result == ""
    mock_post.assert_not_called()


def test_vision_budget_enforces_attempt_and_time_limits():
    attempt_budget = VisionBudget(max_attempts=1)
    assert attempt_budget.try_acquire() is True
    assert attempt_budget.try_acquire() is False

    expired_budget = VisionBudget(max_attempts=5, max_duration_seconds=0)
    assert expired_budget.try_acquire() is False
    assert expired_budget.exhausted is True


@pytest.mark.asyncio
async def test_vision_transcribe_disabled_fallback():
    # When disabled, should immediately fall back to RapidOCR
    img_bytes = _create_test_image(400, 300)
    result = await VisionService.transcribe_image_async(
        img_bytes,
        enabled=False,
    )
    assert isinstance(result, str)


@pytest.mark.asyncio
async def test_vision_transcribe_mocked_success():
    img_bytes = _create_test_image(400, 300)
    
    mock_resp = MagicMock()
    mock_resp.status_code = 200
    mock_resp.json.return_value = {
        "choices": [
            {
                "message": {
                    "content": "| 提单号 | 船名 | 航次 |\n| --- | --- | --- |\n| COSU12345 | CSCL GLOBE | 045E |"
                }
            }
        ]
    }

    with patch("httpx.AsyncClient.post", new_callable=AsyncMock) as mock_post:
        mock_post.return_value = mock_resp
        
        result = await VisionService.transcribe_image_async(
            img_bytes,
            enabled=True,
            custom_api_key="mock_test_key",
            custom_base_url="https://api.sensenova.cn/compatible-mode/v1",
            custom_model="SenseChat-Vision",
        )
        assert "COSU12345" in result
        assert "CSCL GLOBE" in result
        assert mock_post.await_args.kwargs["json"]["max_tokens"] == 4096


def test_vision_transcribe_sync_caps_model_output():
    img_bytes = _create_test_image(400, 300)
    mock_resp = MagicMock()
    mock_resp.status_code = 200
    mock_resp.json.return_value = {
        "choices": [{"message": {"content": "X" * (MAX_VISION_OUTPUT_CHARS + 100)}}]
    }

    with patch("httpx.Client.post", return_value=mock_resp) as mock_post:
        result = VisionService.transcribe_image_sync(
            img_bytes,
            enabled=True,
            custom_api_key="mock_test_key",
        )

    assert len(result) == MAX_VISION_OUTPUT_CHARS
    assert mock_post.call_args.kwargs["json"]["max_tokens"] == 4096


@pytest.mark.asyncio
async def test_vision_transcribe_error_fallback():
    img_bytes = _create_test_image(400, 300)
    
    # Simulate API 500 error
    mock_resp = MagicMock()
    mock_resp.status_code = 500
    mock_resp.text = "Internal Server Error"

    with patch("httpx.AsyncClient.post", new_callable=AsyncMock) as mock_post:
        mock_post.return_value = mock_resp
        
        result = await VisionService.transcribe_image_async(
            img_bytes,
            enabled=True,
            custom_api_key="mock_test_key",
        )
        # Should gracefully fall back to RapidOCR without throwing exception
        assert isinstance(result, str)


def test_word_parser_with_embedded_image(tmp_path):
    # Create Word docx with text + embedded image screenshot
    doc_file = tmp_path / "booking_with_screenshot.docx"
    doc = docx.Document()
    doc.add_paragraph("Please find attached booking details:")
    
    img_bytes = _create_test_image(600, 400)
    img_stream = io.BytesIO(img_bytes)
    doc.add_picture(img_stream, width=docx.shared.Inches(4.0))
    doc.save(str(doc_file))

    with patch.object(VisionService, "transcribe_image_sync", return_value="CONTAINER NO: OOCU9876543 / SEAL: SL12345"):
        text, tables, ocr_text = parse_word(doc_file)
        assert "Please find attached booking details:" in text
        assert "OOCU9876543" in ocr_text


def test_pdf_parser_with_embedded_image(tmp_path):
    # Create a real image-only PDF page (an attachment is not a page image).
    pdf_file = tmp_path / "scanned_doc.pdf"
    Image.new("RGB", (600, 400), color=(255, 255, 255)).save(
        pdf_file, format="PDF", resolution=100
    )

    with patch.object(
        VisionService,
        "transcribe_image_sync",
        return_value="BOOKING: MSK55667788",
    ) as mock_transcribe:
        text, tables, ocr_text = parse_pdf(pdf_file)
        mock_transcribe.assert_called_once()
        assert "MSK55667788" in ocr_text


def test_word_parser_counts_failed_attempts_against_task_budget(tmp_path):
    doc_file = tmp_path / "two_images.docx"
    doc = docx.Document()
    doc.add_picture(io.BytesIO(_create_test_image(600, 400, color=(255, 0, 0))))
    doc.add_picture(io.BytesIO(_create_test_image(600, 400, color=(0, 0, 255))))
    doc.save(doc_file)
    budget = VisionBudget(max_attempts=1)

    with patch.object(VisionService, "transcribe_image_sync", return_value="") as mock_transcribe:
        parse_word(doc_file, budget)

    assert budget.attempts == 1
    assert mock_transcribe.call_count == 1


@pytest.mark.asyncio
async def test_worker_refreshes_persisted_vision_controls():
    await init_db()
    keys = {
        "VISION_LLM_ENABLED": "false",
        "VISION_LLM_MODEL": "vision-worker-test",
        "VISION_LLM_TIMEOUT_SECONDS": "44",
        "VISION_MAX_IMAGES_PER_TASK": "3",
    }
    async with AsyncSessionLocal() as db:
        await db.execute(delete(SystemConfig).where(SystemConfig.key.in_(keys)))
        db.add_all(SystemConfig(key=key, value=value) for key, value in keys.items())
        await db.commit()

    with (
        patch.object(settings, "VISION_LLM_ENABLED", True),
        patch.object(settings, "VISION_LLM_MODEL", "stale-model"),
        patch.object(settings, "VISION_LLM_TIMEOUT_SECONDS", 30),
        patch.object(settings, "VISION_MAX_IMAGES_PER_TASK", 5),
    ):
        await VisionService.refresh_runtime_settings()
        assert settings.VISION_LLM_ENABLED is False
        assert settings.VISION_LLM_MODEL == "vision-worker-test"
        assert settings.VISION_LLM_TIMEOUT_SECONDS == 44
        assert settings.VISION_MAX_IMAGES_PER_TASK == 3

    async with AsyncSessionLocal() as db:
        await db.execute(delete(SystemConfig).where(SystemConfig.key.in_(keys)))
        await db.commit()


def test_parse_single_file_image_dispatch(tmp_path):
    img_file = tmp_path / "cargo_seal.png"
    img_file.write_bytes(_create_test_image(500, 300))

    with patch.object(VisionService, "transcribe_image_sync", return_value="SEAL NO: ML-CN-8899"):
        attachment = parse_single_file(img_file)
        assert attachment.content_type == "image/png"
        assert "ML-CN-8899" in attachment.ocr_text
