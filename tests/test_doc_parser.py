from email.message import EmailMessage
from pathlib import Path
import struct

import docx
import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient

from app.config import Settings, settings
from app.core.parser import parse_single_file, process_uploaded_files
from app.core.parser import doc_parser
from app.core.parser.doc_parser import DocParseError, _extract_word_stream_text, parse_doc
from app.core.parser.eml_parser import parse_eml
from app.database import init_db
from app.main import app


@pytest_asyncio.fixture(autouse=True)
async def prepare_db():
    await init_db()


def _build_word97_streams(
    pieces: list[tuple[str, bool]],
    *,
    codec: str = "cp1252",
    chs: int = 0,
) -> tuple[bytes, bytes]:
    """Build valid FIB and CLX/Piece Table streams without faking an OLE container."""
    word_stream = bytearray(1_024)
    struct.pack_into("<H", word_stream, 0, 0xA5EC)  # wIdent
    struct.pack_into("<H", word_stream, 2, 0x00C1)  # Word 97 nFib
    struct.pack_into("<H", word_stream, 6, 0x0804 if chs == 0x86 else 0x0409)
    struct.pack_into("<H", word_stream, 10, 0x0200 | 0x1000)  # 1Table + Unicode capable
    struct.pack_into("<H", word_stream, 20, chs)
    struct.pack_into("<H", word_stream, 32, 14)  # csw
    struct.pack_into("<H", word_stream, 62, 22)  # cslw
    struct.pack_into("<H", word_stream, 152, 34)  # cbRgFcLcb, includes fcClx

    codepoints = [0]
    pcd_records = []
    file_offset = 1_024
    first_text_offset = file_offset
    for text, compressed in pieces:
        encoded = text.encode(codec if compressed else "utf-16le")
        word_stream.extend(encoded)
        fc_value = file_offset * 2 | 0x40000000 if compressed else file_offset
        pcd_records.append(b"\x00\x00" + struct.pack("<I", fc_value) + b"\x00\x00")
        file_offset += len(encoded)
        # Compressed pieces use one CP per stored byte, including DBCS lead/trail bytes.
        codepoints.append(codepoints[-1] + (len(encoded) if compressed else len(text)))

    struct.pack_into("<I", word_stream, 24, first_text_offset)
    struct.pack_into("<I", word_stream, 28, file_offset)
    struct.pack_into("<I", word_stream, 76, codepoints[-1])  # ccpText

    plc_pcd = b"".join(struct.pack("<I", cp) for cp in codepoints) + b"".join(pcd_records)
    clx = b"\x02" + struct.pack("<I", len(plc_pcd)) + plc_pcd
    struct.pack_into("<II", word_stream, 418, 0, len(clx))  # fcClx/lcbClx
    return bytes(word_stream), clx


def _build_compound_doc(word_stream: bytes, table_stream: bytes) -> bytes:
    """Wrap WordDocument and 1Table in a minimal valid CFB v3 container."""
    sector_size = 512
    stream_size = 4_096
    if len(word_stream) > stream_size or len(table_stream) > stream_size:
        raise ValueError("Test streams exceed the single-chain fixture size")

    free_sector = 0xFFFFFFFF
    end_of_chain = 0xFFFFFFFE
    fat_sector = 0xFFFFFFFD
    no_stream = 0xFFFFFFFF

    header = bytearray(sector_size)
    header[:8] = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
    struct.pack_into("<HHHH", header, 24, 0x003E, 3, 0xFFFE, 9)
    struct.pack_into("<H", header, 32, 6)
    struct.pack_into("<I", header, 40, 0)
    struct.pack_into("<I", header, 44, 1)
    struct.pack_into("<I", header, 48, 0)
    struct.pack_into("<I", header, 56, 4_096)
    struct.pack_into("<I", header, 60, end_of_chain)
    struct.pack_into("<I", header, 68, end_of_chain)
    for index in range(109):
        struct.pack_into("<I", header, 76 + index * 4, free_sector)
    struct.pack_into("<I", header, 76, 17)  # FAT sector

    def directory_entry(
        name: str,
        entry_type: int,
        left: int,
        right: int,
        child: int,
        first_sector: int,
        size: int,
    ) -> bytes:
        entry = bytearray(128)
        encoded_name = name.encode("utf-16le") + b"\x00\x00"
        entry[:len(encoded_name)] = encoded_name
        struct.pack_into("<HBBIII", entry, 64, len(encoded_name), entry_type, 1, left, right, child)
        struct.pack_into("<I", entry, 116, first_sector)
        struct.pack_into("<Q", entry, 120, size)
        return bytes(entry)

    directory = b"".join(
        [
            directory_entry("Root Entry", 5, no_stream, no_stream, 1, end_of_chain, 0),
            directory_entry("WordDocument", 2, 2, no_stream, no_stream, 1, stream_size),
            directory_entry("1Table", 2, no_stream, no_stream, no_stream, 9, stream_size),
            bytes(128),
        ]
    )

    fat_entries = [free_sector] * (sector_size // 4)
    fat_entries[0] = end_of_chain
    for sector in range(1, 8):
        fat_entries[sector] = sector + 1
    fat_entries[8] = end_of_chain
    for sector in range(9, 16):
        fat_entries[sector] = sector + 1
    fat_entries[16] = end_of_chain
    fat_entries[17] = fat_sector
    fat = b"".join(struct.pack("<I", value) for value in fat_entries)

    return b"".join(
        [
            bytes(header),
            directory,
            word_stream.ljust(stream_size, b"\x00"),
            table_stream.ljust(stream_size, b"\x00"),
            fat,
        ]
    )


def test_parse_doc_rejects_empty_corrupt_and_propagates_failure(tmp_path):
    empty_file = tmp_path / "empty.doc"
    empty_file.write_bytes(b"")
    with pytest.raises(DocParseError, match="empty"):
        parse_doc(empty_file)

    corrupt_file = tmp_path / "corrupt.doc"
    corrupt_file.write_bytes(b"RANDOM_GARBAGE_BYTES_1234567890")
    with pytest.raises(DocParseError, match="not a supported"):
        parse_doc(corrupt_file)
    with pytest.raises(DocParseError):
        parse_single_file(corrupt_file)
    with pytest.raises(DocParseError):
        process_uploaded_files([corrupt_file], temp_dir=tmp_path)


def test_parse_doc_misnamed_docx(tmp_path):
    document = docx.Document()
    document.add_heading("BOOKING CONFIRMATION", level=1)
    document.add_paragraph("Carrier: COSCO SHIPPING")
    document.add_paragraph("Booking No: COSU12345678")
    doc_path = tmp_path / "misnamed.doc"
    document.save(str(doc_path))

    text, tables, ocr = parse_doc(doc_path)
    assert "BOOKING CONFIRMATION" in text
    assert "COSU12345678" in text
    assert "COSCO SHIPPING" in text
    assert ocr == ""


def test_parse_doc_rtf_unicode_codepage_and_destinations(tmp_path):
    rtf_content = (
        br"{\rtf1\ansi\ansicpg936\uc1 {\fonttbl {\f0 Courier;}}"
        br"\f0\fs24 SHIPPER: ABC LOGISTICS\par Unicode: \u20320?\u22909?"
        br"\par Hex: \'c4\'e3\'ba\'c3\par POD: ROTTERDAM\par "
        br"FIELD\cell VALUE\cell\row POL\cell NINGBO\cell\row}"
    )
    rtf_path = tmp_path / "rtf_booking.doc"
    rtf_path.write_bytes(rtf_content)

    text, tables, ocr = parse_doc(rtf_path)
    assert "ABC LOGISTICS" in text
    assert "ROTTERDAM" in text
    assert text.count("你好") == 2
    assert "Courier" not in text
    assert tables == [
        {"table_index": 0, "rows": [["FIELD", "VALUE"], ["POL", "NINGBO"]]}
    ]
    assert ocr == ""


def test_parse_doc_html_decodes_entities_and_ignores_script(tmp_path):
    html_content = (
        "<html><body><h1>BOOKING &amp; ADVICE</h1>"
        "<script>SECRET_SCRIPT_TEXT</script><table><tr><td>POL</td><td>SHANGHAI</td></tr></table>"
        "</body></html>"
    )
    html_path = tmp_path / "advice.doc"
    html_path.write_text(html_content, encoding="utf-8")

    text, tables, ocr = parse_doc(html_path)
    assert "BOOKING & ADVICE" in text
    assert "SHANGHAI" in text
    assert "SECRET_SCRIPT_TEXT" not in text
    assert tables == [{"table_index": 0, "rows": [["POL", "SHANGHAI"]]}]


def test_word97_piece_table_extracts_unicode_and_compressed_text():
    word_stream, table_stream = _build_word97_streams(
        [("订舱号：MSCU9876543\r", False), ("POL: SHANGHAI\rPOD: HAMBURG", True)],
    )

    text = _extract_word_stream_text(word_stream, table_stream)
    assert "订舱号：MSCU9876543" in text
    assert "POL: SHANGHAI" in text
    assert "POD: HAMBURG" in text
    assert text.index("MSCU9876543") < text.index("POL: SHANGHAI")

    chinese_word_stream, chinese_table_stream = _build_word97_streams(
        [("中文压缩文本", True)],
        codec="cp936",
        chs=0x86,
    )
    assert _extract_word_stream_text(chinese_word_stream, chinese_table_stream) == "中文压缩文本"


def test_parse_doc_uses_valid_compound_file_and_word_streams(tmp_path):
    word_stream, table_stream = _build_word97_streams(
        [("BOOKING: ONEY1234567\rFIELD\x07VALUE\x07\rPOL\x07NINGBO\x07\r", False)]
    )
    doc_path = tmp_path / "legacy_booking.doc"
    doc_path.write_bytes(_build_compound_doc(word_stream, table_stream))

    assert doc_parser.olefile.isOleFile(str(doc_path))
    text, tables, ocr = parse_doc(doc_path)
    assert "ONEY1234567" in text
    assert tables == [
        {"table_index": 0, "rows": [["FIELD", "VALUE"], ["POL", "NINGBO"]]}
    ]
    assert ocr == ""


def test_parse_doc_rejects_fake_ole_and_encrypted_fib(tmp_path):
    fake_ole_path = tmp_path / "fake_ole.doc"
    fake_ole_path.write_bytes(
        b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 504 + "BOOKING".encode("utf-16le")
    )
    with pytest.raises(DocParseError, match="Invalid Word 97-2003 OLE"):
        parse_doc(fake_ole_path)

    word_stream, table_stream = _build_word97_streams([("SECRET", False)])
    encrypted_stream = bytearray(word_stream)
    flags = struct.unpack_from("<H", encrypted_stream, 10)[0]
    struct.pack_into("<H", encrypted_stream, 10, flags | 0x0100)
    with pytest.raises(DocParseError, match="Encrypted"):
        _extract_word_stream_text(bytes(encrypted_stream), table_stream)


def test_parse_doc_wraps_signature_probe_errors(monkeypatch, tmp_path):
    doc_path = tmp_path / "probe_error.doc"
    doc_path.write_bytes(b"not a recognized document")

    def fail_probe(_):
        raise OSError("probe failed")

    monkeypatch.setattr(doc_parser.olefile, "isOleFile", fail_probe)
    with pytest.raises(DocParseError, match="parse the Word document safely"):
        parse_single_file(doc_path)


def test_parse_doc_enforces_specific_size_limit(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "MAX_LEGACY_DOC_FILE_SIZE", 1_024)
    oversized = tmp_path / "oversized.doc"
    oversized.write_bytes(b"X" * 1_025)
    with pytest.raises(DocParseError, match="size limit"):
        parse_doc(oversized)


def test_legacy_doc_limit_cannot_exceed_generic_upload_limit():
    with pytest.raises(ValueError, match="MAX_LEGACY_DOC_FILE_SIZE"):
        Settings(
            _env_file=None,
            MAX_UPLOAD_FILE_SIZE=1_024,
            MAX_LEGACY_DOC_FILE_SIZE=2_048,
            MAX_UPLOAD_TOTAL_SIZE=4_096,
        )


def test_parse_single_file_doc_dispatch(tmp_path):
    document = docx.Document()
    document.add_paragraph("SHIPPER: GLOBAL FREIGHT LTD")
    document.add_paragraph("CONTAINER: TGHU1234567 40HQ")
    doc_path = tmp_path / "test_booking.doc"
    document.save(str(doc_path))

    attachment = parse_single_file(doc_path)
    assert attachment.filename == "test_booking.doc"
    assert attachment.content_type == "application/msword"
    assert "GLOBAL FREIGHT LTD" in attachment.text
    assert "TGHU1234567" in attachment.text


def test_eml_parser_with_doc_attachment(tmp_path):
    eml_file = tmp_path / "booking_with_doc.eml"
    doc_file = tmp_path / "attachment.doc"
    document = docx.Document()
    document.add_paragraph("Booking Ref: MSK889900")
    document.save(str(doc_file))

    message = EmailMessage()
    message["Subject"] = "Fwd: Booking Advice"
    message["From"] = "carrier@shipping.com"
    message["To"] = "agent@cargoplus.com"
    message.set_content("Please find attached booking document in doc format.")
    message.add_attachment(
        doc_file.read_bytes(),
        maintype="application",
        subtype="msword",
        filename="booking_advice.doc",
    )
    eml_file.write_bytes(message.as_bytes())

    output_dir = tmp_path / "extracted_att"
    subject, body, extracted_files = parse_eml(eml_file, output_dir)
    assert "Booking Advice" in subject
    assert len(extracted_files) == 1
    assert "MSK889900" in parse_single_file(extracted_files[0]).text


@pytest.mark.asyncio
async def test_extract_async_upload_with_doc_and_specific_limit(tmp_path, monkeypatch):
    transport = ASGITransport(app=app)

    async with AsyncClient(transport=transport, base_url="http://test") as client:
        register_response = await client.post(
            "/api/v1/auth/register",
            json={
                "company_name": f"DocTester_{Path(tmp_path).name[:8]}",
                "contact_email": f"doc_{Path(tmp_path).name[:8]}@example.com",
                "password": "Password123!",
            },
        )
        assert register_response.status_code == 200
        tenant_id = register_response.json()["data"]["tenant_id"]
        api_key = register_response.json()["data"]["api_key"]
        await client.put(
            f"/admin/tenants/{tenant_id}/status?is_active=true",
            headers={"X-Admin-Secret": settings.ADMIN_SECRET_KEY},
        )

        document = docx.Document()
        document.add_paragraph("BOOKING NO: ONEY1234567")
        document.add_paragraph("VESSEL: ONE APUS / 012E")
        document.add_paragraph("POL: NINGBO")
        document.add_paragraph("POD: LONG BEACH")
        sample_doc = tmp_path / "booking_note.doc"
        document.save(str(sample_doc))

        with sample_doc.open("rb") as handle:
            upload_response = await client.post(
                "/api/v1/extract/async/upload",
                headers={"Authorization": f"Bearer {api_key}"},
                files=[("files", ("booking_note.doc", handle, "application/msword"))],
                data={"mail_subject": "Booking advice .doc upload"},
            )
        assert upload_response.status_code == 200
        assert "task_id" in upload_response.json()

        monkeypatch.setattr(settings, "MAX_LEGACY_DOC_FILE_SIZE", 1_024)
        with sample_doc.open("rb") as handle:
            oversized_response = await client.post(
                "/api/v1/extract/async/upload",
                headers={"Authorization": f"Bearer {api_key}"},
                files=[("files", ("booking_note.doc", handle, "application/msword"))],
            )
        assert oversized_response.status_code == 413
        assert oversized_response.json()["detail"]["code"] == 41301
