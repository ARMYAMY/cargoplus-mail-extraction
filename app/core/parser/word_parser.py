import logging
from pathlib import Path
from typing import Any, List, Tuple
from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from app.services.vision_service import VisionBudget, VisionService

logger = logging.getLogger(__name__)


def _table_rows(table: Table) -> List[List[str]]:
    rows = []
    for row in table.rows[:200]:
        row_data = [cell.text.strip()[:2_000] for cell in row.cells[:100]]
        if any(row_data):
            rows.append(row_data)
    return rows


def _table_markdown(rows: List[List[str]]) -> str:
    """Serialize a Word table without dropping empty cells or column positions."""
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    padded_rows = [row + [""] * (width - len(row)) for row in rows]
    lines = [
        "| " + " | ".join(padded_rows[0]) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in padded_rows[1:])
    return "\n".join(lines)


def parse_word(
    file_path: Path,
    vision_budget: VisionBudget | None = None,
) -> Tuple[str, List[Any], str]:
    """
    Parses a Word (.docx) document.
    Extracts text, structured tables, and transcribes embedded document screenshots/images via Vision LLM.
    Returns: (text, tables, ocr_text)
    """
    paragraphs = []
    tables = []
    ocr_parts = []

    try:
        doc = Document(str(file_path))

        # python-docx exposes paragraphs and tables in separate collections,
        # but shipping instructions rely on their original interleaved order.
        # Use the document body order for real documents and retain the legacy
        # collection fallback for lightweight mocks and older integrations.
        if isinstance(doc, DocumentObject):
            body_items = (
                ("paragraph" if isinstance(item, Paragraph) else "table", item)
                for item in doc.iter_inner_content()
                if isinstance(item, (Paragraph, Table))
            )
        else:  # pragma: no cover - exercised by mocked compatibility tests
            body_items = [
                *(("paragraph", item) for item in doc.paragraphs[:5_000]),
                *(("table", item) for item in doc.tables[:20]),
            ]

        table_index = 0
        paragraph_count = 0
        for item_kind, item in body_items:
            if item_kind == "paragraph":
                if paragraph_count >= 5_000:
                    continue
                paragraph_count += 1
                if item.text.strip():
                    paragraphs.append(item.text.strip()[:5_000])
                continue

            if item_kind != "table" or table_index >= 20:
                continue
            rows = _table_rows(item)
            if rows:
                tables.append({"table_index": table_index, "rows": rows})
                paragraphs.append(_table_markdown(rows))
            table_index += 1

        # Extract embedded images / screenshots from Word package
        if vision_budget is None:
            from app.config import settings

            vision_budget = VisionBudget(settings.VISION_MAX_IMAGES_PER_TASK)
        if hasattr(doc, "part") and hasattr(doc.part, "related_parts"):
            for rel_id, rel_part in doc.part.related_parts.items():
                if vision_budget.exhausted:
                    logger.info("Reached task-wide image transcription limit in Word %s", file_path.name)
                    break

                content_type = getattr(rel_part, "content_type", "").lower()
                part_name = str(getattr(rel_part, "partname", "")).lower()
                if "image" in content_type or "image" in part_name:
                    img_bytes = getattr(rel_part, "blob", None)
                    if img_bytes and VisionService.is_valid_document_image(img_bytes):
                        if not vision_budget.try_acquire():
                            break
                        try:
                            img_text = VisionService.transcribe_image_sync(
                                img_bytes,
                                filename_hint=f"{file_path.name}_embedded_{rel_id}",
                                custom_timeout=vision_budget.request_timeout(),
                            )
                            if img_text.strip():
                                ocr_parts.append(
                                    f"[Word文档内嵌单证截图/图片识别内容]:\n{img_text.strip()}"
                                )
                        except Exception as img_err:
                            logger.warning(f"Error transcribing Word embedded image in {file_path.name}: {img_err}")

    except Exception as e:
        logger.error(f"Failed to parse Word document {file_path}: {e}")
        return "", [], f"Error parsing Word: {str(e)}"

    full_text = "\n\n".join(paragraphs)
    full_ocr = "\n\n".join(ocr_parts)

    return full_text, tables, full_ocr
